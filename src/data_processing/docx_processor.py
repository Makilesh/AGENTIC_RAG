"""
DOCX Processor for Agentic RAG System.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..utils import get_data_processing_logger
from .semantic_chunker import Chunk, SemanticChunker

logger = get_data_processing_logger()


@dataclass
class DOCXSection:
    """
    Represents a section in a DOCX document.
    
    Attributes:
        heading: Section heading text.
        heading_level: Heading level (1-9, or 0 for body text).
        content: Text content of the section.
        tables: Tables within this section.
    """
    heading: str
    heading_level: int
    content: str
    tables: List[str] = field(default_factory=list)


@dataclass
class DOCXDocument:
    """
    Represents a processed DOCX document.
    
    Attributes:
        file_name: Name of the DOCX file.
        file_path: Full path to the file.
        sections: List of document sections.
        full_text: Combined text from all sections.
        metadata: Document-level metadata.
    """
    file_name: str
    file_path: str
    sections: List[DOCXSection]
    full_text: str
    metadata: dict = field(default_factory=dict)


class DOCXProcessor:
    """
    Process DOCX documents for RAG ingestion.
    
    Extracts text with style preservation, maintains heading hierarchy,
    and processes embedded tables.
    """
    
    # Mapping of Word heading styles to levels
    HEADING_STYLES = {
        'Heading 1': 1,
        'Heading 2': 2,
        'Heading 3': 3,
        'Heading 4': 4,
        'Heading 5': 5,
        'Heading 6': 6,
        'Heading 7': 7,
        'Heading 8': 8,
        'Heading 9': 9,
        'Title': 0,
    }
    
    def __init__(self, chunker: Optional[SemanticChunker] = None):
        """
        Initialize the DOCX processor.
        
        Args:
            chunker: SemanticChunker instance. If None, creates a new one.
        """
        self.chunker = chunker or SemanticChunker()
        logger.info("DOCXProcessor initialized")
    
    def _get_heading_level(self, paragraph: Paragraph) -> int:
        """
        Determine the heading level of a paragraph.
        
        Args:
            paragraph: python-docx Paragraph object.
            
        Returns:
            Heading level (1-9) or 0 if not a heading.
        """
        style_name = paragraph.style.name if paragraph.style else ""
        
        # Check for explicit heading styles
        for style, level in self.HEADING_STYLES.items():
            if style.lower() in style_name.lower():
                return level
        
        # Check for outline level in style
        try:
            if paragraph.style.paragraph_format.outline_level is not None:
                return paragraph.style.paragraph_format.outline_level + 1
        except AttributeError:
            pass
        
        return 0
    
    def _table_to_markdown(self, table: Table) -> str:
        """
        Convert a Word table to markdown format.
        
        Args:
            table: python-docx Table object.
            
        Returns:
            Markdown formatted table string.
        """
        rows = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.replace("|", "\\|").strip()
                cell_text = re.sub(r'\s+', ' ', cell_text)
                cells.append(cell_text)
            rows.append(cells)
        
        if not rows:
            return ""
        
        # Build markdown
        lines = []
        
        # Header row
        header = "| " + " | ".join(rows[0]) + " |"
        lines.append(header)
        
        # Separator
        separator = "|" + "|".join(["---"] * len(rows[0])) + "|"
        lines.append(separator)
        
        # Data rows
        for row in rows[1:]:
            # Pad row to match header length
            while len(row) < len(rows[0]):
                row.append("")
            row_str = "| " + " | ".join(row[:len(rows[0])]) + " |"
            lines.append(row_str)
        
        return "\n".join(lines)
    
    def _iter_block_items(self, parent) -> list:
        """
        Generate a list of paragraphs and tables in document order.
        
        Args:
            parent: Document or cell to iterate over.
            
        Yields:
            Paragraph or Table objects in order.
        """
        if isinstance(parent, DocumentType):
            parent_elm = parent.element.body
        else:
            parent_elm = parent._element
        
        items = []
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                items.append(Paragraph(child, parent))
            elif isinstance(child, CT_Tbl):
                items.append(Table(child, parent))
        
        return items
    
    def _extract_text_with_formatting(self, paragraph: Paragraph) -> str:
        """
        Extract text with basic formatting indicators.
        
        Args:
            paragraph: python-docx Paragraph object.
            
        Returns:
            Text with formatting context.
        """
        text_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Add formatting context
            if run.bold and run.italic:
                text_parts.append(f"***{text}***")
            elif run.bold:
                text_parts.append(f"**{text}**")
            elif run.italic:
                text_parts.append(f"*{text}*")
            else:
                text_parts.append(text)
        
        return "".join(text_parts)
    
    def extract(self, file_path: str) -> DOCXDocument:
        """
        Extract content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            DOCXDocument containing extracted content.
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        if not path.suffix.lower() == '.docx':
            raise ValueError(f"Not a DOCX file: {file_path}")
        
        logger.info(f"Processing DOCX: {path.name}")
        
        doc = Document(file_path)
        sections = []
        full_text_parts = []
        
        current_section = DOCXSection(
            heading="Document Start",
            heading_level=0,
            content=""
        )
        current_content_parts = []
        current_tables = []
        
        for item in self._iter_block_items(doc):
            if isinstance(item, Paragraph):
                heading_level = self._get_heading_level(item)
                text = self._extract_text_with_formatting(item)
                
                if heading_level > 0 and text.strip():
                    # Save current section
                    if current_content_parts or current_tables:
                        current_section.content = "\n".join(current_content_parts)
                        current_section.tables = current_tables
                        sections.append(current_section)
                        full_text_parts.append(current_section.content)
                    
                    # Start new section
                    current_section = DOCXSection(
                        heading=text.strip(),
                        heading_level=heading_level,
                        content=""
                    )
                    current_content_parts = []
                    current_tables = []
                else:
                    if text.strip():
                        current_content_parts.append(text)
                        
            elif isinstance(item, Table):
                table_md = self._table_to_markdown(item)
                if table_md:
                    current_tables.append(table_md)
                    current_content_parts.append(f"\n[TABLE]\n{table_md}\n[/TABLE]\n")
        
        # Save last section
        if current_content_parts or current_tables:
            current_section.content = "\n".join(current_content_parts)
            current_section.tables = current_tables
            sections.append(current_section)
            full_text_parts.append(current_section.content)
        
        # Combine full text
        full_text = "\n\n".join(full_text_parts)
        
        # Build metadata
        metadata = {
            "source_type": "docx",
            "file_name": path.name,
            "file_path": str(path.absolute()),
            "num_sections": len(sections),
            "has_tables": any(s.tables for s in sections)
        }
        
        logger.info(
            f"Extracted {len(sections)} sections from {path.name}",
        )
        
        return DOCXDocument(
            file_name=path.name,
            file_path=str(path.absolute()),
            sections=sections,
            full_text=full_text,
            metadata=metadata
        )
    
    def process(self, file_path: str) -> List[Chunk]:
        """
        Extract and chunk a DOCX document.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            List of Chunk objects ready for embedding.
        """
        # Extract content
        docx_doc = self.extract(file_path)
        
        # Prepare base metadata
        base_metadata = {
            "source_type": "docx",
            "file_name": docx_doc.file_name,
            "file_path": docx_doc.file_path,
            "has_tables": docx_doc.metadata.get("has_tables", False)
        }
        
        all_chunks = []
        
        # Process section by section
        for section in docx_doc.sections:
            if not section.content.strip():
                continue
            
            section_metadata = {
                **base_metadata,
                "heading_level": section.heading_level,
                "section_title": section.heading
            }
            
            # Chunk the section content
            section_chunks = self.chunker.chunk(
                section.content,
                f"{docx_doc.file_name}_{section.heading[:30]}",
                section_metadata
            )
            all_chunks.extend(section_chunks)
        
        # Re-index all chunks
        for i, chunk in enumerate(all_chunks):
            chunk.chunk_index = i
            chunk.chunk_id = f"{docx_doc.file_name}_chunk_{i}"
            chunk.metadata["chunk_index"] = i
            chunk.metadata["total_chunks"] = len(all_chunks)
        
        logger.info(
            f"Created {len(all_chunks)} chunks from DOCX {docx_doc.file_name}",
        )
        
        return all_chunks
